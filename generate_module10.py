import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_module10_proposal():
    doc = docx.Document()

    # Set normal font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_title_page():
        for _ in range(4):
            doc.add_paragraph()
        
        paragraphs = [
            "Module 10 Assignment: Strategic Vendor Decision Support System",
            "Final Project Report and Dashboard Implementation",
            "Sumesh Chakkaravarthi Purushothaman",
            "Northeastern University",
            "ALY 6980: Capstone",
            "Dr. Valerie Atherley"
        ]
        
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            if text in paragraphs[:2]:
                for run in p.runs:
                    run.bold = True
        
        doc.add_page_break()

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            run.font.bold = True

    def add_paragraph(text, indent=True):
        p = doc.add_paragraph(text)
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_reference(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    # Building document
    add_title_page()

    add_heading("Overview and Rationale", level=1)
    add_paragraph("The Summit Global Consulting project aims to create a unified decision-making system using multiple vendor and partner data sources by transforming their separate data into a single strategic decision support system for business operations. The initial business problem which involves merging multiple spreadsheets into one standard system requires operational solutions but the actual work for the project requires them to solve data governance problems and maintain data accuracy while using ethical analytics methods to assess their healthcare service procurement needs and their partnerships with healthcare providers.")
    add_paragraph("As healthcare organizations increasingly work with third-party vendors to deliver their clinical and operational functions many organizations experience difficulties because they cannot access identical vendor data which exists in the United States through various vendor-specific sources which present vendor data in different formats and assess service quality to vendors through compliance-based methods instead of value-based methods. The healthcare executive team faces challenges which make it difficult to reach fair and responsible vendor use decisions because of vendor-related data assessment difficulties. Summit Global Consulting enables organizations to gather all vendor data into one database which allows them to exit their present method of gathering vendor data without using it for ongoing vendor assessment. The organization will have an active and structured approach to create essential information about vendor data which enables them to establish logical and defendable vendor management decisions through their vendor assessment process which includes both supplier diversity and operational efficiency evaluations.")

    add_heading("Purpose", level=2)
    add_paragraph("The goal of this initiative is to shorten the time to get products to the market by helping Summit Global Consulting locate and rank high-value IT-industry vendors that 1) demonstrate a committed supplier diversity and that 2) can be verified to have current, active operations and that 3) can be easily contacted. This will be accomplished by moving from passive reporting to structured vendor intelligence and converting raw data into reasoned decisions.", indent=False)

    add_heading("Audience", level=2)
    add_paragraph("This system is designed specifically for business development managers, procurement analysts, supplier diversity officers and executive leadership. These groups will expect high-confidence analytics that can be easily understood, can be defended and align with their strategic outreach objectives. Therefore, this system has been developed for non-technical business professionals rather than data scientists.", indent=False)

    add_heading("Scope", level=2)
    add_paragraph("This initiative is purposefully limited to: IT-sector vendors, validated vendor contact information, active vendor operations, and the top 20-50 ranked vendors based upon various criteria. Also, predictive modelling and external benchmarking will not be used so as to avoid bias and maintain transparency in the governance and methodology employed.", indent=False)

    add_heading("Intersection with Data Governance, Ethics, and Social Responsibility", level=1)
    add_paragraph("To advance beyond initial proposals, this deployed system aggressively tackles data governance through standardized schemas and transparent accountability measures. Vendor datasets previously lacked uniformity, hindering the organization’s ability to assess performance and diversity commitments against strategic objectives. By enforcing a rigorous definition of data elements—particularly mandating verified contact information—the system establishes crucial data integrity and traceability. Ethically, the framework elevates corporate social responsibility by converting supplier diversity from a symbolic checklist into a core performance metric. As executive leaders interact with the deployed dashboards, they rely on analytics that support, rather than replace, human judgment, ensuring that procurement strategies remain unbiased and operationally equitable.")

    add_heading("Literature Review", level=1)
    
    add_heading("1. Adobor & McMullen (2014)", level=2)
    add_paragraph("Adobor and McMullen (2014) demonstrate that supplier diversity programs offer substantial business value beyond mere compliance. By integrating diverse suppliers into their networks, organizations gain critical flexibility to adapt to market fluctuations. The authors argue that supplier diversity metrics should function as strategic assets, particularly in complex industries. This research directly supports our dashboard methodology, which shifts the paradigm from passively reporting Supplier Diversity Office (SDO) metrics to actively using them as empirical criteria for ranking and prioritizing vendor relationships.")

    add_heading("2. Whitfield & Landeros (2006)", level=2)
    add_paragraph("In their investigation of supply chain inclusivity, Whitfield and Landeros (2006) revealed that major organizations consistently block diverse suppliers due to information asymmetry, the absence of transparent directories, and poor contact management. They concluded that establishing centralized, highly organized vendor information architectures is the primary remedy. Our project translates this finding into action: the entire foundation of the proposed analytics platform is designed to eliminate these transparency barriers by mandating current contact verification, ensuring decision-makers can instantly connect with diverse IT suppliers.")

    add_heading("3. Obermeyer & Emanuel (2016)", level=2)
    add_paragraph("Obermeyer and Emanuel (2016) warned of the ethical implications surrounding the rapid expansion of predictive analytics and machine learning in healthcare. They asserted that while advanced models offer distinct advantages, they must operate under strict, unbiased parameters and leave ultimate decision-making to human judgment. Our initiative explicitly adopts this philosophy. By intentionally avoiding opaque predictive modeling and instead providing highly interpretable, deterministic dashboard analytics, the project guarantees that organizational leaders maintain executive accountability and avoid systemic algorithmic bias.")

    add_heading("4. Khatri & Brown (2010)", level=2)
    add_paragraph("Khatri and Brown (2010) established that data governance is a comprehensive framework—comprising people, processes, and practices—necessary for delivering reliable data in heavily regulated sectors. Their central thesis is that without an effective governance structure, technology alone cannot salvage analytical output. This directly informs the core logic of our system, where data trustworthiness takes priority over data completeness; vendors lacking verified attributes are systematically filtered out, ensuring executive stakeholders base procurement strategies only on demonstrably factual intelligence.")

    add_heading("5. Davenport & Harris (2007)", level=2)
    add_paragraph("Davenport and Harris (2007) explored the stages of analytics maturity, concluding that true competitive advantage requires integrating analytics directly into everyday operational processes. They found that highly complex models often fail if they do not align with user workflows. Influenced by this research, Summit Global Consulting explicitly prioritized usability. By designing intuitive, interactive dashboards tailored for non-technical stakeholders, the project bridges the gap between raw data collection and strategic execution without requiring advanced data science expertise.")

    add_heading("Advanced Exploratory Data Analysis Integration", level=1)
    add_paragraph("To realize the vision outlined in the literature, a substantial Exploratory Data Analysis (EDA) phase was implemented to dissect IT vendor concentrations and SDO commitments. The analysis transformed raw spreadsheets into high-quality visual insights that map vendor distribution across Massachusetts agencies and highlight critical relationships between diversity percentages and total business spend. This graphical intelligence calibrated the logic required for the scoring algorithm, uncovering underlying data biases before dashboard initialization.")
    
    img_path_1 = '/Users/sumesh/Projects/Antigravity/Capstone/output/BQ3_Vendor_Distribution_Treemap.png'
    if os.path.exists(img_path_1):
        try:
            doc.add_picture(img_path_1, width=Inches(6.0))
            p_cap = doc.add_paragraph("Figure 1. Vendor Distribution Treemap indicating category concentrations.", indent=False)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            pass

    add_heading("Interactive Dashboard Deployment", level=1)
    add_paragraph("Fulfilling the fundamental purpose established in the proposal, an interactive Plotly Dash dashboard was engineered to serve as the unified strategic decision support system. The dashboard consumes the standardized, governance-cleared dataset, projecting it into a multi-tabbed interface featuring dynamic Key Performance Indicators (KPIs) and actionable vendor rankings. Business Development Managers can now intuitively filter IT-sector companies by total commitments, assessing the market landscape instantaneously. This bridges the transition from administrative data collection to robust, visually guided operational intelligence.")

    img_path_2 = '/Users/sumesh/Projects/Antigravity/Capstone/dashboard_screenshot.png'
    if os.path.exists(img_path_2):
        try:
            doc.add_picture(img_path_2, width=Inches(6.0))
            p_cap2 = doc.add_paragraph("Figure 2. Strategic Vendor Dashboard Interface.", indent=False)
            p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            pass

    add_heading("Conclusion", level=1)
    add_paragraph("This Module 10 assignment represents the culmination of a rigorous transition from a theoretical proposal into a tangible analytics solution. By anchoring the development solidly in the core objectives of creating a unified, ethically sound, and interpretable vendor system, the final implementation successfully mitigates the initial challenges of spreadsheet fragmentation. Through the application of dynamic dashboarding and rigorous data governance rules, Summit Global Consulting is empowered to execute socially responsible, highly efficient procurement initiatives without the risks associated with opaque predictive algorithms.")

    doc.add_page_break()

    add_heading("References", level=1)
    refs = [
        "Adobor, H., & McMullen, R. S. (2014). Strategic purchasing and supplier partnerships—The role of a supplier diversity program. Journal of Business & Industrial Marketing, 29(4), 263–272. https://doi.org/10.1108/JBIM-03-2012-0051",
        "Davenport, T. H., & Harris, J. G. (2007). Competing on analytics: The new science of winning. Harvard Business School Press.",
        "Khatri, V., & Brown, C. V. (2010). Designing data governance. Communications of the ACM, 53(1), 148–152. https://doi.org/10.1145/1629175.1629210",
        "Obermeyer, Z., & Emanuel, E. J. (2016). Predicting the future—Big data, machine learning, and clinical medicine. The New England Journal of Medicine, 375(13), 1216–1219. https://doi.org/10.1056/NEJMp1606181",
        "Whitfield, G., & Landeros, R. (2006). Supplier diversity: Portals to prosperity? Supply Chain Management Review, 10(5), 16–23. https://www.scmr.com/article/supplier_diversity_portals_to_prosperity"
    ]
    for r in refs:
        add_reference(r)

    doc.save('/Users/sumesh/Projects/Antigravity/Capstone/ALY6980_Module10_Strong_Report.docx')
    print("Module 10 Word document created successfully at /Users/sumesh/Projects/Antigravity/Capstone/ALY6980_Module10_Strong_Report.docx")

if __name__ == "__main__":
    create_module10_proposal()
