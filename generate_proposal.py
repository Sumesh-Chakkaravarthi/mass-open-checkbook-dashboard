import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_proposal():
    doc = docx.Document()

    # Set normal font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Helper functions
    def add_title_page():
        for _ in range(3):
            doc.add_paragraph()
        
        paragraphs = [
            "Individual Project Proposal: Strategic Vendor Decision Support System",
            "Sumesh Chakkaravarthi Purushothaman",
            "Northeastern University",
            "ALY 6980: Capstone",
            "Dr. Valerie Atherley",
            "February 17, 2026"
        ]
        
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            
            # Make title bold
            if text == paragraphs[0]:
                for run in p.runs:
                    run.bold = True
        
        doc.add_page_break()

    def add_heading(text, level=1):
        # Docx built-in headings have different default styles, we override to keep Times New Roman and black
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            run.font.bold = True

    def add_paragraph(text, indent=True):
        p = doc.add_paragraph(text)
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_reference(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    # Building document
    add_title_page()

    add_heading("Individual Project Proposal: Strategic Vendor Decision Support System", level=1)
    
    add_heading("Introduction", level=1)
    
    add_heading("Statement of the Problem", level=2)
    add_paragraph("Healthcare organizations are increasingly dependent on third-party vendors for critical clinical and operational functions. A major challenge executive teams face is the fragmented nature of vendor data. Vendor information is dispersed across multiple spreadsheets and external databases, lacking standardized formats. Furthermore, vendors are often evaluated purely on compliance metrics rather than value-based indicators. This scattered and unstandardized data creates barriers to making fair, responsible, and timely vendor selection decisions, ultimately resulting in lengthy procurement cycles and market entry delays.")

    add_heading("Purpose of the Study", level=2)
    add_paragraph("The primary business purpose of this project is to accelerate time-to-market for healthcare operations by developing a unified decision-making system. This initiative aims to identify, evaluate, and rank high-value information technology (IT) sector vendors based on verifiable operational status, contactability, and a demonstrated commitment to supplier diversity. By transforming raw, disorganized data into structured operational intelligence, Summit Global Consulting can transition from passive vendor reporting to active, strategic vendor management.")

    add_heading("Review of the Literature", level=2)
    add_paragraph("Five key areas of literature inform the rationale and methodology of this proposal. Adobor and McMullen (2014) emphasized that integrating a supplier diversity program builds supply chain resilience and flexibility. Their research underscores the importance of developing standard methodologies for categorizing and ranking supplier firms based on performance metrics, rather than simply tracking diversity quotas. Whitfield and Landeros (2006) expanded on the barriers minority and diverse suppliers face, specifically information asymmetry and the lack of transparent organizational contact points. They argued that centralized vendor information systems are essential for improving purchaser access to qualified diverse vendors, which validates the need for a curated, transparent analytics directory.")
    add_paragraph("The ethical deployment of analytical systems is another critical component. Obermeyer and Emanuel (2016) highlighted the rapid expansion of predictive analytics in clinical and operational settings, cautioning that these systems must be unbiased and auditable. They asserted that analytics should support, rather than replace, human judgment. This aligns directly with the goal of creating transparent decision support dashboards instead of automated black-box predictive models. Establishing reliable data for these dashboards requires robust governance. Khatri and Brown (2010) defined data governance as a framework necessary for delivering high-quality, accountable data. They noted that analytical outputs are functionally useless without strict data trustworthiness—a principle that justifies excluding unverified vendors from the proposed model. Finally, Davenport and Harris (2007) provided a framework for analytics maturity, arguing that successful implementation requires integrating analytical tools directly into existing operational workflows. Their work supports designing simple, accessible dashboard interfaces tailored for non-technical business professionals to ensure actual organizational adoption.")

    add_heading("Approach", level=2)
    add_paragraph("The approach utilizes a dashboard-centric analytics model. Rather than employing complex predictive modeling, the project addresses the core business need through targeted data governance, schema standardization, and interactive visualization. Disconnected spreadsheets will be consolidated into a governed database, which will feed into an executive dashboard. This tool will automatically filter IT vendors and apply a weighted ranking logic to present the most viable, diverse, and contactable partners.")

    add_heading("Methods", level=1)
    
    add_heading("Data Collection Instruments, Variables, and Materials", level=2)
    add_paragraph("Data collection leverages a mix of internal proprietary materials and public databases. Internal instruments include existing organizational vendor spreadsheets containing historical procurement records. External materials include public procurement data from USAspending.gov, the Massachusetts state procurement portal, and Supplier Diversity Office (SDO) directories. Specific variables extracted for analysis include Vendor ID, primary Industry/Sector, Service Code, SDO diversity commitment percentage, current contract status, historical performance measures, and verified geographic service areas.")

    add_heading("Considerations on Data Governance", level=2)
    add_paragraph("Strict data governance is the foundation of this proposal. The data ingestion process establishes a standardized schema to harmonize classification categories across disparate sources. A crucial governance rule implemented is the requirement for verified active contact information; vendors failing this verification check are excluded from the final database to preserve the integrity and usability of the dashboard. This ensures the executive leadership team bases decisions exclusively on reliable, actionable data.")

    add_heading("Data Analysis Procedures", level=2)
    add_paragraph("The analytical procedure begins with data cleaning, designed to remove duplicate records and standardize formatting. Exploratory data analysis evaluates the distribution of key parameters, such as SDO commitment percentages and regional clustering. This exploratory phase informs the calibration of the vendor scoring model. Following exploration, a weighted evaluation algorithm processes the curated data. The proposed weighting framework allocates 40% to SDO Commitment Percentage, 30% to Normalized Addressable Spend, 20% to Active Contract Status, and 10% to Contact Verification. The resulting scores are integrated into an interactive dashboard, allowing users to filter and sort vendors dynamically.")

    add_heading("Anticipated Ethical Issues", level=2)
    add_paragraph("Several ethical risks require mitigation. There is a potential to overweight diversity metrics at the expense of operational capability, which could jeopardize project feasibility. Additionally, strict data verification criteria might inadvertently introduce exclusion bias against smaller vendors lacking robust public footprints. Finally, there may be stakeholder resistance to an algorithmic vendor ranking system. To mitigate these issues, the scoring model features adjustable weighting parameters, and all rules for vendor exclusion are thoroughly documented and transparent. The system functions strictly as an advisory tool, leaving final authority with human leadership.")

    add_heading("Preliminary Studies or Pilot Tests", level=2)
    add_paragraph("Preliminary validation involves descriptive analysis and internal stakeholder review. Initial exploratory data analysis acts as a pilot test of the raw datasets to identify structural biases or outliers before locking in the scoring weights. Additionally, an internal stakeholder validation survey will be deployed to business development and procurement staff. This survey blends Likert-scale questions with open-ended feedback to assess the perceived reliability of the selected data variables and the acceptability of the proposed weighting logic. Survey results will guide iterative refinements to the dashboard before final deployment.")

    add_heading("Conclusion and Recommendations", level=1)
    add_paragraph("The proposed vendor prioritization system provides a practical mechanism to embed governance, ethics, and social responsibility directly into the purchasing lifecycle. By replacing disconnected spreadsheets with a structured, transparent analytical platform, the organization can substantially reduce vendor onboarding delays while advancing corporate diversity goals. It is recommended that the project advance to the data collection and cleaning phase, prioritizing the establishment of the standardized schema to maintain the proposed 10 to 12-week delivery timeline.")

    doc.add_page_break()

    add_heading("References", level=1)
    refs = [
        "Adobor, H., & McMullen, R. S. (2014). Strategic purchasing and supplier partnerships—The role of a supplier diversity program. Journal of Business & Industrial Marketing, 29(4), 263–272. https://doi.org/10.1108/JBIM-03-2012-0051",
        "Davenport, T. H., & Harris, J. G. (2007). Competing on analytics: The new science of winning. Harvard Business School Press.",
        "Khatri, V., & Brown, C. V. (2010). Designing data governance. Communications of the ACM, 53(1), 148–152. https://doi.org/10.1145/1629175.1629210",
        "Obermeyer, Z., & Emanuel, E. J. (2016). Predicting the future—Big data, machine learning, and clinical medicine. The New England Journal of Medicine, 375(13), 1216–1219. https://doi.org/10.1056/NEJMp1606181",
        "Whitfield, G., & Landeros, R. (2006). Supplier diversity: Portals to prosperity? Supply Chain Management Review, 10(5), 16–23. https://www.scmr.com/article/supplier_diversity_portals_to_prosperity"
    ]
    for r in refs:
        add_reference(r)

    doc.add_page_break()

    add_heading("Appendices", level=1)
    add_heading("Appendix A: Proposed Project Timeline", level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Duration'
    
    rows = [
        ('Data Collection & Enrichment', '3–4 weeks'),
        ('Cleaning & Standardization', '2 weeks'),
        ('Exploratory Data Analysis & Weight Calibration', '2 weeks'),
        ('Dashboard Development', '3 weeks'),
        ('Survey Validation & Refinement', '2 weeks'),
        ('Total Estimated Timeline', '10–12 weeks')
    ]
    for phase, duration in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = duration

    doc.save('/Users/sumesh/Projects/Antigravity/Capstone/Individual_Project_Proposal_Draft.docx')
    print("Word document created successfully at /Users/sumesh/Projects/Antigravity/Capstone/Individual_Project_Proposal_Draft.docx")

if __name__ == "__main__":
    create_proposal()
