#!/usr/bin/env python3
"""
Generate the EDA Report as a Word Document (.docx)
Includes all business questions with embedded plots.
Also generates BQ11: Sector Analysis (Company Density by Industry).
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import seaborn as sns
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import textwrap
import warnings

warnings.filterwarnings('ignore')

# ── Config ──
VENDOR_FILE = "Copy of Vendor Contact Details (1).xlsx"
CATEGORIZED_FILE = "List of Categorized_Companies (1).xlsx"
OUTPUT_DIR = Path("output")
SKIP_SHEETS = ['Abbreviations ']

CATEGORY_NAMES = {
    'ITE': 'IT Equipment & Services',
    'ITS': 'IT Software & Services',
    'ITT': 'Telecom & Networking',
    'FAC': 'Facilities General',
    'VEH': 'Vehicle Acquisition & Maint.',
    'GRO': 'Food & Food Service',
    'LND': 'Facility Landscaping',
    'MED': 'Health & Medical',
    'MRO': 'Maintenance, Repair & Ops',
    'OFF': 'Office Supplies',
    'PRF': 'Professional Services',
    'PSE': 'Public Safety & Security',
    'SFC': 'Sustainable Facilities',
    'TRD': 'Tradespersons',
    'WMR': 'Waste Mgmt & Recycling',
}

plt.rcParams.update({
    'figure.facecolor': '#FAFAFA',
    'axes.facecolor': '#FAFAFA',
    'axes.edgecolor': '#CCCCCC',
    'axes.labelcolor': '#333333',
    'xtick.color': '#555555',
    'ytick.color': '#555555',
    'font.family': 'sans-serif',
    'font.size': 11,
    'axes.titlesize': 14,
    'axes.titleweight': 'bold',
})


# ── Data Loading (reuse from eda_analysis.py) ──
def load_vendor_data(filepath):
    xl = pd.ExcelFile(filepath)
    frames = []
    for sheet in xl.sheet_names:
        if sheet.strip() in [s.strip() for s in SKIP_SHEETS]:
            continue
        df = pd.read_excel(filepath, sheet_name=sheet, header=0).iloc[:, :7]
        df.columns = ['Contract_Code', 'Name', 'Company', 'Role', 'Email', 'Phone', 'SDO_Pct']
        df['Category'] = sheet.strip()
        metadata_kw = ['Master Contract', 'Solicitation Enabled', 'Master MBPO',
                        'Bid and Contract', 'Category and Vendor', 'Category Development',
                        'Mass Gov', 'OSD Help Desk', 'N/A']
        mask = df['Company'].astype(str).apply(
            lambda x: not any(kw.lower() in x.lower() for kw in metadata_kw))
        df = df[mask].copy()
        df['SDO_Pct'] = pd.to_numeric(df['SDO_Pct'], errors='coerce')
        df['Company'] = df['Company'].astype(str).str.replace('\n', ' ', regex=False).str.strip()
        df['Name'] = df['Name'].astype(str).str.replace('\n', ' ', regex=False).str.strip()
        df['Contract_Code'] = df['Contract_Code'].astype(str).str.replace('\n', ' ', regex=False).str.strip()
        frames.append(df)
    return pd.concat(frames, ignore_index=True)


def load_categorized_companies(filepath):
    xl = pd.ExcelFile(filepath)
    frames = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(filepath, sheet_name=sheet, header=None)
        industry = sheet.strip()
        for col_idx, typ in [(1, 'National & Local'), (2, 'Local'), (3, 'SGC Target')]:
            companies = df.iloc[2:, col_idx].dropna().astype(str).str.strip().tolist()
            for c in companies:
                if c and c != 'nan':
                    frames.append({'Industry': industry, 'Company': c, 'Type': typ})
    return pd.DataFrame(frames)


def get_label(code):
    return CATEGORY_NAMES.get(code, code)


# ── Generate BQ11: Sector Analysis ──
def generate_sector_analysis(cat_df):
    """BQ11: Company density by industry sector."""
    industry_counts = cat_df.groupby('Industry')['Company'].nunique().sort_values(ascending=True)

    fig, ax = plt.subplots(figsize=(12, 7))
    colors = sns.color_palette("mako", len(industry_counts))
    bars = ax.barh(range(len(industry_counts)), industry_counts.values, color=colors, edgecolor='white', linewidth=0.5)

    ax.set_yticks(range(len(industry_counts)))
    ax.set_yticklabels(industry_counts.index, fontsize=10)
    ax.set_xlabel('Number of Unique Companies', fontweight='bold')
    ax.set_title('BQ11: Company Density by Industry Sector', pad=15, fontsize=14, fontweight='bold')

    for bar, val in zip(bars, industry_counts.values):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                str(val), va='center', fontsize=10, fontweight='bold', color='#333')

    fig.tight_layout()
    path = OUTPUT_DIR / 'BQ11_Sector_Company_Density.png'
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    print(f"   📊 Saved: {path.name}")
    return path


# ── Build Word Document ──
def create_report(vendor_df, cat_df):
    doc = Document()

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Title Page ──
    doc.add_paragraph('')
    doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Massachusetts Open Checkbook')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Vendor Contract & SDO Commitment Analysis')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph('')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('ALY 6980 — Capstone Project\nExploratory Data Analysis Report\n\nFebruary 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Dashboard link
    dash_para = doc.add_paragraph()
    dash_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dash_para.add_run('Interactive Dashboard: ')
    run.font.size = Pt(11)
    run.font.bold = True
    run = dash_para.add_run('http://127.0.0.1:8050')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)
    run.font.underline = True

    doc.add_page_break()

    # ── Table of Contents (manual) ──
    toc_title = doc.add_heading('Table of Contents', level=1)
    business_questions = [
        ("BQ1", "Top IT Sector Companies by SDO Commitment %"),
        ("BQ2", "Average SDO Commitment % by Procurement Category"),
        ("BQ3", "Distribution of Vendors Across Procurement Categories"),
        ("BQ4", "Top Contract Sub-Categories by Vendor Count"),
        ("BQ5", "National vs Local Company Presence by Industry"),
        ("BQ6", "SDO Commitment Distribution & Outlier Detection"),
        ("BQ7", "SDO Coverage Rate — Valid vs Missing"),
        ("BQ8", "Correlation: Vendor Count vs Average SDO"),
        ("BQ9", "Industry Diversity — National vs Local vs SGC Target"),
        ("BQ10", "IT Sector Vendor Concentration"),
        ("BQ11", "Sector Analysis: Company Density by Industry"),
    ]
    for bq_id, bq_title in business_questions:
        p = doc.add_paragraph(f'{bq_id}: {bq_title}', style='List Number')

    doc.add_page_break()

    # ── Business Question Pages ──
    # Pre-compute stats
    sdo_df = vendor_df.dropna(subset=['SDO_Pct'])
    sdo_df = sdo_df[sdo_df['SDO_Pct'] > 0]
    sdo_capped = sdo_df.copy()
    sdo_capped['SDO_Capped'] = sdo_capped['SDO_Pct'].clip(upper=1.0)

    it_df = vendor_df[vendor_df['Category'].isin(['ITE', 'ITS', 'ITT'])]
    it_sdo = it_df.dropna(subset=['SDO_Pct'])
    it_sdo = it_sdo[it_sdo['SDO_Pct'] > 0]

    bq_details = [
        {
            'id': 'BQ1',
            'title': 'Which companies in the IT sector (ITE, ITS, ITT) have the best SDO commitment percentage?',
            'plot': 'BQ1_IT_Top_SDO_Companies.png',
            'analysis': (
                f"Among the three IT sub-categories, {len(it_sdo)} vendor records have valid SDO commitments. "
                f"NEWCOM Wireless Services leads with 51% SDO commitment (ITE category), followed by "
                f"Digit Outsource Inc at 46% (ITS) and CenturyLink at 36% (ITS). "
                f"The average SDO commitment across IT vendors is {it_sdo['SDO_Pct'].mean():.1%}, "
                f"with a median of {it_sdo['SDO_Pct'].median():.1%}, indicating most IT companies "
                f"have relatively modest SDO commitments with a few standout performers."
            ),
        },
        {
            'id': 'BQ2',
            'title': 'Which industry/category has the highest average SDO commitment percentage?',
            'plot': 'BQ2_Avg_SDO_by_Category.png',
            'analysis': (
                f"Across all 15 procurement categories, the average SDO commitment varies significantly. "
                f"Categories like Professional Services and Facility Landscaping show higher average commitments, "
                f"while Food & Food Service and Vehicle Acquisition show lower averages. "
                f"This disparity highlights that supplier diversity policies are not uniformly adopted "
                f"across all procurement sectors, suggesting targeted policy interventions may be needed."
            ),
        },
        {
            'id': 'BQ3',
            'title': 'What is the distribution of vendors across procurement categories?',
            'plot': 'BQ3_Vendor_Distribution_Treemap.png',
            'analysis': (
                f"The treemap reveals significant variation in vendor density across categories. "
                f"Professional Services (PRF) has the largest vendor pool, reflecting the broad "
                f"range of consulting and advisory services Massachusetts procures. "
                f"IT Equipment and Facilities General also have substantial vendor bases. "
                f"Categories like Food & Food Service and Tradespersons have fewer vendors, "
                f"suggesting more concentrated markets."
            ),
        },
        {
            'id': 'BQ4',
            'title': 'How are vendors distributed across contract sub-category codes?',
            'plot': 'BQ4_Contract_Subcategory_Vendors.png',
            'analysis': (
                "The top contract sub-categories by vendor count are PRF76 (116 vendors), "
                "PRF74 (109 vendors), and PSE01 (67 vendors). Professional Services dominates "
                "with multiple high-vendor-count sub-categories, indicating a fragmented and "
                "competitive market. This granular view helps procurement officers identify "
                "which specific contract areas have the most supplier options."
            ),
        },
        {
            'id': 'BQ5',
            'title': 'How do National vs Local companies compare across industries?',
            'plot': 'BQ5_National_vs_Local_Industries.png',
            'analysis': (
                "Finance & Insurance has the most National & Local companies (27), followed by "
                "Technology & Engineering (24) and Healthcare & Biotechnology (21). "
                "The SGC (Supplier/Government/Community) targets are relatively consistent across "
                "major industries at around 8-9 companies each. Hospitality & Entertainment has "
                "the fewest presence overall, likely reflecting Massachusetts' economic profile "
                "where finance, tech, and healthcare dominate."
            ),
        },
        {
            'id': 'BQ6',
            'title': 'What is the SDO commitment distribution? Are there outliers?',
            'plot': 'BQ6_SDO_Distribution_Boxplot.png',
            'analysis': (
                "The box plot reveals significant outliers across multiple categories. "
                "Facilities General and Sustainable Facilities show the most extreme outliers "
                "with some vendors reporting SDO commitments above 100% (likely data quality issues). "
                "Most categories have median SDO commitments between 1-5%, with IQR-based outlier "
                "detection identifying 80+ outliers across 12 of 15 categories. "
                "This suggests a need for data validation in SDO reporting."
            ),
        },
        {
            'id': 'BQ7',
            'title': 'What proportion of vendors have valid SDO commitments?',
            'plot': 'BQ7_SDO_Coverage_Rate.png',
            'analysis': (
                "SDO coverage varies dramatically by category. Facility Landscaping leads with 93.2% "
                "of vendors reporting SDO commitments, while IT categories lag significantly: "
                "ITE at 13.5%, ITS at 7.4%, and ITT at 6.5%. Tradespersons has 0% coverage. "
                "This stark difference suggests that SDO reporting requirements may not be uniformly "
                "enforced across all procurement categories, particularly in the IT sector."
            ),
        },
        {
            'id': 'BQ8',
            'title': 'Is there a correlation between vendor count and average SDO commitment?',
            'plot': 'BQ8_Vendor_Count_vs_SDO_Scatter.png',
            'analysis': (
                "The Pearson correlation coefficient is r = 0.395 (p = 0.162), indicating a weak "
                "positive but statistically non-significant relationship between the number of "
                "vendors in a category and the average SDO commitment. This suggests that market "
                "competition alone does not drive supplier diversity — policy mandates and category-"
                "specific requirements are likely more influential factors."
            ),
        },
        {
            'id': 'BQ9',
            'title': 'How diverse is the industry representation among categorized companies?',
            'plot': 'BQ9_Industry_Diversity_Heatmap.png',
            'analysis': (
                "The heatmap shows that Finance & Insurance and Healthcare & Biotechnology have "
                "the most diverse company representation across all three types (National & Local, "
                "Local, SGC Target). Technology & Engineering has a strong National & Local presence "
                "(24 companies) but relatively fewer Local-only companies (5), suggesting the tech "
                "sector is dominated by national firms. The SGC target gaps are most visible in "
                "Government & Non-Profit (only 3 targets vs 15 national companies)."
            ),
        },
        {
            'id': 'BQ10',
            'title': 'Do a few companies dominate the IT sector? (Vendor Concentration)',
            'plot': 'BQ10_IT_Vendor_Concentration_Donut.png',
            'analysis': (
                f"The IT sector has {it_df['Company'].nunique()} unique companies across ITE, ITS, and ITT. "
                f"The top 10 companies hold only about 8.5% of total contract associations, indicating "
                f"a highly fragmented market with no single dominant vendor. This healthy competition "
                f"suggests that Massachusetts has successfully avoided vendor lock-in in IT procurement, "
                f"providing leverage for better pricing and service quality negotiations."
            ),
        },
        {
            'id': 'BQ11',
            'title': 'Sector Analysis: Company Density by Industry',
            'plot': 'BQ11_Sector_Company_Density.png',
            'analysis': (
                "This sector analysis visualizes the company density across 10 industry classifications "
                "in the Massachusetts categorized companies dataset. Finance & Insurance leads with the "
                "highest density, followed by Healthcare & Biotechnology and Technology & Engineering. "
                "Hospitality & Entertainment and Publishing & Media have the lowest density, reflecting "
                "Massachusetts' economic specialization in financial services, life sciences, and "
                "technology sectors. This insight is valuable for policymakers targeting industry-specific "
                "supplier diversity programs."
            ),
        },
    ]

    for i, bq in enumerate(bq_details):
        # Heading
        heading = doc.add_heading(f"{bq['id']}: {bq['title']}", level=2)

        # Plot
        plot_path = OUTPUT_DIR / bq['plot']
        if plot_path.exists():
            doc.add_picture(str(plot_path), width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[Plot not found: {bq['plot']}]")

        # Analysis
        doc.add_paragraph('')
        analysis_heading = doc.add_paragraph()
        run = analysis_heading.add_run('Analysis & Key Findings:')
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)

        analysis_para = doc.add_paragraph(bq['analysis'])
        analysis_para.paragraph_format.space_after = Pt(6)

        # Page break after each (except last)
        if i < len(bq_details) - 1:
            doc.add_page_break()

    # ── Final page: Dashboard Link ──
    doc.add_page_break()
    doc.add_heading('Interactive Dashboard', level=1)
    p = doc.add_paragraph('The interactive dashboard provides additional features beyond the static EDA plots:')
    doc.add_paragraph('KPI summary cards (Total Vendors, Average SDO, Categories)', style='List Bullet')
    doc.add_paragraph('Interactive IT sub-category filter (ITE / ITS / ITT / All)', style='List Bullet')
    doc.add_paragraph('Radar chart comparing IT sub-categories', style='List Bullet')
    doc.add_paragraph('SDO histogram across all vendors', style='List Bullet')
    doc.add_paragraph('Hover tooltips and zoom on all charts', style='List Bullet')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Dashboard URL: ')
    run.font.bold = True
    run = p.add_run('http://127.0.0.1:8050')
    run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)
    run.font.underline = True

    p = doc.add_paragraph()
    run = p.add_run('To run: ')
    run.font.bold = True
    run = p.add_run('python dashboard.py')
    run.font.name = 'Courier New'

    # Save
    output_path = Path('ALY6980_EDA_Report.docx')
    doc.save(str(output_path))
    print(f"\n✅ Report saved: {output_path.resolve()}")
    return output_path


# ── Main ──
if __name__ == '__main__':
    print("Loading data...")
    vendor_df = load_vendor_data(VENDOR_FILE)
    cat_df = load_categorized_companies(CATEGORIZED_FILE)
    print(f"✅ {len(vendor_df):,} vendors, {len(cat_df):,} categorized companies")

    # Generate BQ11 plot
    print("\nGenerating BQ11: Sector Analysis...")
    generate_sector_analysis(cat_df)

    # Create Word report
    print("\nBuilding Word document...")
    create_report(vendor_df, cat_df)
